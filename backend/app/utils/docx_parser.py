"""
DOCX 解析器 (DOCX Parser)

从 .docx 文件中提取段落文本和格式信息。

解析流程:
  1. 使用 python-docx 打开 .docx 文件
  2. 遍历所有段落，跳过空段落
  3. 对每个非空段落提取:
     - 文本内容 (text)
     - 格式信息 (style_info): 字体名、字号、加粗、斜体、下划线、对齐方式、行距、首行缩进
     - 是否为标题及标题级别

格式提取策略:
  - 段落级格式 (对齐、行距、缩进): 从 paragraph.paragraph_format 提取
  - 字符级格式 (字体、字号、加粗、斜体、下划线): 从段落的第一个 run 提取
  - 注意: 同一段落内混合格式（部分加粗/部分普通）在导出时会简化

标题识别:
  检查段落的样式名称是否以 "Heading" 开头
  - "Heading 1" → heading_level=1
  - "Heading 2" → heading_level=2
  以此类推

对齐映射:
  WD_ALIGN_PARAGRAPH.LEFT    → 0
  WD_ALIGN_PARAGRAPH.CENTER  → 1
  WD_ALIGN_PARAGRAPH.RIGHT   → 2
  WD_ALIGN_PARAGRAPH.JUSTIFY → 3

扩展指南:
  - 如需提取颜色、背景色等格式，在 extract_style_info() 中添加对应字段
  - 如需支持表格内容提取，新增 extract_tables() 函数
  - 对齐映射值如需调整，只需修改 ALIGNMENT_MAP 字典
"""

import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.schemas.document import StyleInfo

logger = logging.getLogger(__name__)

# python-docx 对齐枚举 → 我们定义的数字编码
# 编码设计: 0=左, 1=中, 2=右, 3=两端对齐
ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: 0,
    WD_ALIGN_PARAGRAPH.CENTER: 1,
    WD_ALIGN_PARAGRAPH.RIGHT: 2,
    WD_ALIGN_PARAGRAPH.JUSTIFY: 3,
}


def parse_docx(file_path: str | Path) -> list[dict]:
    """
    解析 .docx 文件，提取所有非空段落及其格式信息

    参数:
        file_path: .docx 文件的路径
    返回:
        list[dict]: 每个元素包含:
          - index (int): 段落序号（从 0 开始）
          - text (str): 段落文本内容
          - style_info (dict): 格式信息字典（由 StyleInfo.model_dump() 生成）
    注意:
        空段落（仅包含空白字符）会被自动跳过，不计入段落总数。
    """
    doc = DocxDocument(str(file_path))
    paragraphs = []

    index = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue  # 跳过空段落

        # 提取该段落的格式信息
        style_info = extract_style_info(para)
        paragraphs.append({
            "index": index,
            "text": text,
            "style_info": style_info.model_dump(),
        })
        index += 1

    logger.info(f"Parsed {len(paragraphs)} non-empty paragraphs from {file_path}")
    return paragraphs


def extract_docx_title(file_path: str | Path) -> str | None:
    """
    尝试从 .docx 文件中提取文档标题

    提取策略（按优先级）:
      1. 文档核心属性中的 title 字段
      2. 文档中第一个 Heading 样式的非空段落文本

    参数:
        file_path: .docx 文件的路径
    返回:
        标题文本，如果两者都没有则返回 None
    """
    doc = DocxDocument(str(file_path))

    # 策略 1: 从文档属性中获取标题（用户在 Word 中设置的"标题"属性）
    if doc.core_properties.title:
        return doc.core_properties.title

    # 策略 2: 取第一个标题样式的段落文本
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip():
            return para.text.strip()

    return None


def extract_style_info(para) -> StyleInfo:
    """
    从 python-docx 段落对象中提取格式化信息

    提取的内容:
      段落级:
        - alignment: 对齐方式（编码为数字 0-3）
        - line_spacing: 行距值
        - first_line_indent: 首行缩进（单位: pt）
        - is_heading / heading_level: 是否为标题及其级别

      字符级（取自第一个 run）:
        - font_name: 字体名称（如 "Times New Roman"）
        - font_size: 字号（单位: pt）
        - bold: 是否加粗
        - italic: 是否斜体
        - underline: 是否下划线

    参数:
        para: python-docx 的 Paragraph 对象
    返回:
        StyleInfo Pydantic 模型
    """
    # ── 段落级格式 ──

    # 对齐方式
    alignment = None
    if para.alignment is not None:
        alignment = ALIGNMENT_MAP.get(para.alignment)

    # 行距
    line_spacing = None
    if para.paragraph_format.line_spacing:
        line_spacing = float(para.paragraph_format.line_spacing)

    # 首行缩进 (转换为 pt)
    first_line_indent = None
    if para.paragraph_format.first_line_indent:
        first_line_indent = para.paragraph_format.first_line_indent.pt

    # 标题检测
    is_heading = para.style.name.startswith("Heading") if para.style else False
    heading_level = None
    if is_heading:
        try:
            # 从样式名称提取级别数字: "Heading 1" → 1, "Heading 2" → 2
            heading_level = int(para.style.name.replace("Heading", "").strip())
        except ValueError:
            heading_level = 1  # 无法解析时默认为一级标题

    # ── 字符级格式（取自第一个 run）──
    font_name = None
    font_size = None
    bold = False
    italic = False
    underline = False

    if para.runs:
        first_run = para.runs[0]
        if first_run.font.name:
            font_name = first_run.font.name
        if first_run.font.size:
            font_size = first_run.font.size.pt
        bold = bool(first_run.bold)
        italic = bool(first_run.italic)
        underline = bool(first_run.underline)

    return StyleInfo(
        font_name=font_name,
        font_size=font_size,
        bold=bold,
        italic=italic,
        underline=underline,
        alignment=alignment,
        line_spacing=line_spacing,
        first_line_indent=first_line_indent,
        is_heading=is_heading,
        heading_level=heading_level,
    )
