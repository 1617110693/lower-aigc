"""
DOCX 构建器 (DOCX Builder)

将降重后的段落数据重新构建为 .docx 文件，保留原始格式。

构建流程:
  1. 创建空白 python-docx Document
  2. 设置默认字体（Times New Roman + 宋体作为东亚字体）
  3. 遍历段落数据，根据 style_info 恢复格式
  4. 将文档写入 BytesIO 缓冲区并返回

格式恢复:
  - 标题段落: 使用 add_heading() 添加（会应用内置标题样式）
  - 普通段落: 创建 Paragraph，依次设置对齐、行距、首行缩进
  - 字符格式: 在 Run 上设置字体名、字号、加粗、斜体、下划线
  - 东亚字体: 同时设置 w:eastAsia 属性，确保中日韩文字使用正确字体

注意:
  - 同一段落内的混合格式（部分加粗/部分普通）在导出时会统一为该段落的
    第一个 run 格式（与解析阶段一致）
  - 标题级别限制在 1-9 级（python-docx 支持的范围）

扩展指南:
  - 如需添加页眉/页脚/页码，在 build_docx() 中操作 doc.sections[0]
  - 如需自定义样式模板，可以在创建 Document 时传入模板路径
"""

import io
import logging

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

from app.schemas.document import StyleInfo

logger = logging.getLogger(__name__)


def _set_run_east_asian_font(run, font_name: str) -> None:
    """
    安全地为 Run 设置东亚字体（中日韩文字回退字体）

    参数:
        run: python-docx Run 对象
        font_name: 字体名称（如 "宋体"、"SimSun" 等）
    """
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_name)


def _set_style_east_asian_font(style, font_name: str) -> None:
    """
    安全地为 Style 设置东亚字体（中日韩文字回退字体）

    参数:
        style: python-docx Style 对象
        font_name: 字体名称（如 "宋体"、"SimSun" 等）
    """
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_name)

# 数字编码 → python-docx 对齐枚举（与 docx_parser.ALIGNMENT_MAP 相反）
ALIGNMENT_REVERSE_MAP = {
    0: WD_ALIGN_PARAGRAPH.LEFT,
    1: WD_ALIGN_PARAGRAPH.CENTER,
    2: WD_ALIGN_PARAGRAPH.RIGHT,
    3: WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def build_docx(paragraphs: list[dict]) -> io.BytesIO:
    """
    根据段落数据构建 .docx 文件

    参数:
        paragraphs: 段落数据列表，每个元素包含:
          - text (str): 段落文本（优先使用 reduced_text）
          - style_info (dict | None): 格式信息（由 docx_parser 提取）

    返回:
        io.BytesIO: 构建好的 .docx 文件字节流（已 seek 到开头）
    """
    doc = DocxDocument()

    # 设置默认字体
    # Times New Roman 用于西文，宋体用于东亚文字（中文、日文、韩文）
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    # 安全地设置东亚字体（使用辅助函数，避免直接访问可能不存在的XML元素）
    _set_style_east_asian_font(style, "宋体")

    for para_data in paragraphs:
        text = para_data.get("text", "")
        style_info_raw = para_data.get("style_info")

        if style_info_raw:
            # 有格式信息 — 按格式构建段落
            style_info = StyleInfo.model_validate(style_info_raw)
            _add_paragraph_with_style(doc, text, style_info)
        else:
            # 无格式信息 — 使用默认格式添加
            doc.add_paragraph(text)

    # 将文档写入内存缓冲区
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)  # 重置指针到开头，供 StreamingResponse 读取
    logger.debug(f"Built DOCX with {len(paragraphs)} paragraphs")
    return buf


def _add_paragraph_with_style(doc: DocxDocument, text: str, style: StyleInfo):
    """
    按指定格式向文档中添加一个段落

    处理逻辑:
      - 如果是标题: 使用 doc.add_heading() 添加（级别限制 1-9）
      - 如果是普通段落: 创建段落 → 设置段落格式 → 添加文本 Run → 设置字符格式

    参数:
        doc: python-docx Document 对象
        text: 段落文本内容
        style: StyleInfo 对象，包含完整的格式信息
    """
    # ── 标题处理 ──
    if style.is_heading and style.heading_level:
        # python-docx 的 add_heading 最大支持 9 级标题
        level = min(style.heading_level, 9)
        doc.add_heading(text, level=level)
        return

    # ── 普通段落处理 ──
    para = doc.add_paragraph()

    # 段落对齐
    if style.alignment is not None and style.alignment in ALIGNMENT_REVERSE_MAP:
        para.alignment = ALIGNMENT_REVERSE_MAP[style.alignment]

    # 行距
    if style.line_spacing:
        para.paragraph_format.line_spacing = style.line_spacing

    # 首行缩进
    if style.first_line_indent:
        para.paragraph_format.first_line_indent = Pt(style.first_line_indent)

    # ── 字符格式 ──
    run = para.add_run(text)
    if style.font_name:
        run.font.name = style.font_name
        # 同时设置东亚字体为相同字体（安全访问XML元素）
        _set_run_east_asian_font(run, style.font_name)
    if style.font_size:
        run.font.size = Pt(style.font_size)
    run.bold = style.bold
    run.italic = style.italic
    run.underline = style.underline
